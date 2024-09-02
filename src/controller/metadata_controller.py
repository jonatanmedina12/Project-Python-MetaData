import re
from abc import ABC, abstractmethod
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from PIL import Image
import mimetypes
import docx
from datetime import datetime


class MetaDataConfig(ABC):
    @abstractmethod
    def extract(self, filepath):
        pass


class ImageMetaDataExtractor(MetaDataConfig):
    def extract(self, filepath):
        with Image.open(filepath) as img:
            if img.format in ['JPG', 'JPEG', 'PNG']:
                metadata = img.info
                if img.format in ['JPG', 'JPEG']:
                    exif = img._getexif()
                    if exif:
                        metadata.update({
                            Image.ExifTags.TAGS.get(key, key): value
                            for key, value in exif.items() if key in Image.ExifTags.TAGS
                        })
                return metadata if metadata else {"Error": "No metadata found."}
            else:
                return {"Error": "Unsupported image format."}


class PdfMetadataExtractor(MetaDataConfig):
    def __init__(self):
        self.email_regex = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

    def extract(self, filepath):
        metadata = {}
        with open(filepath, 'rb') as f:
            parser = PDFParser(f)
            doc = PDFDocument(parser)
            if doc.info:
                for info in doc.info:
                    for key, value in info.items():
                        if isinstance(value, bytes):
                            try:
                                decoded_value = value.decode('utf-16be')
                            except UnicodeDecodeError:
                                decoded_value = value.decode('utf-8', errors='ignore')
                        else:
                            decoded_value = value
                        metadata[key] = decoded_value

        text = extract_text(filepath)
        emails = self.extract_email(text)
        metadata['extracted_emails'] = emails
        metadata['text'] = text[:1000]  # First 1000 characters of text

        return metadata

    def extract_email(self, text):
        return re.findall(self.email_regex, text)


class DocxMetadataExtractor(MetaDataConfig):
    def extract(self, filepath):
        doc = docx.Document(filepath)
        prop = doc.core_properties
        attributes = [
            "author", "category", "comments", "content_status",
            "created", "identifier", "keywords", "last_modified_by",
            "language", "modified", "subject", "title", "version"
        ]
        metadata = {attr: getattr(prop, attr, None) for attr in attributes}

        # Convert datetime objects to strings
        for key, value in metadata.items():
            if isinstance(value, datetime):
                metadata[key] = value.isoformat()

        # Extract text content (first 1000 characters)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        metadata['text'] = ' '.join(full_text)[:1000]

        return metadata


class MetadataExtractorFactory:
    @staticmethod
    def get_extractor(filepath):
        mime_type, _ = mimetypes.guess_type(filepath)
        if mime_type:
            if mime_type.startswith('image'):
                return ImageMetaDataExtractor()
            elif mime_type == 'application/pdf':
                return PdfMetadataExtractor()
            elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                return DocxMetadataExtractor()
        raise ValueError(f"Unsupported file type: {mime_type}")


def extract_metadata(filepath):
    try:
        extractor = MetadataExtractorFactory.get_extractor(filepath)
        metadata = extractor.extract(filepath)
        return {"success": True, "metadata": metadata}
    except Exception as e:
        return {"success": False, "error": str(e)}
